from __future__ import annotations

from collections import Counter
from copy import deepcopy
from dataclasses import dataclass
from math import ceil
from pathlib import Path
import hashlib
import re
import shutil
import subprocess
import tempfile
from typing import Iterable, Sequence

from loguru import logger


SUPPORTED_SUFFIXES = {".pdf", ".docx"}
MAX_PAGES_PER_REQUEST = 299
DOCX_WORDS_PER_PAGE = 350
DOCX_SAFE_MAX_PAGES_WITHOUT_EXACT_COUNT = 250
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_P_TAG = f"{{{WORD_NS}}}p"
W_T_TAG = f"{{{WORD_NS}}}t"
W_SECTPR_TAG = f"{{{WORD_NS}}}sectPr"

_WORD_PATTERN = re.compile(r"\w+", re.UNICODE)
_DOCX_TOC_LINE_PATTERN = re.compile(
    r"^(?P<title>.+?)(?:[\t\s\.·…\-]{2,}|\t)(?P<page>\d+)\s*$"
)
_PDF_HEADING_PATTERNS = [
    re.compile(r"^Chapter\s+([A-Z]\d+)\s+(.+)", re.MULTILINE),
    re.compile(r"^Chapter\s+(\d+)[.:]?\s*(.*)", re.MULTILINE),
    re.compile(r"^(\d+)\.\s+([A-Z][A-Za-z\s&\-]+)", re.MULTILINE),
]
_PDF_PAGE_CLONE_IGNORE_FIELDS = ("/Annots", "/Parent", "/B", "/StructParents")


def _load_docx_document_factory():
    try:
        from docx import Document as docx_document_factory
    except Exception as exc:
        raise RuntimeError(
            "DOCX support requires `python-docx` to be installed."
        ) from exc
    return docx_document_factory


def _load_pypdf_classes():
    try:
        from pypdf import PdfReader as pdf_reader, PdfWriter as pdf_writer
    except Exception:
        try:
            from PyPDF2 import PdfReader as pdf_reader, PdfWriter as pdf_writer  # type: ignore
        except Exception as exc:
            raise RuntimeError(
                "PDF support requires `pypdf` (or fallback `PyPDF2`) to be installed."
            ) from exc
    return pdf_reader, pdf_writer


@dataclass(frozen=True)
class BookmarkEntry:
    level: int
    title: str
    page_index: int


@dataclass(frozen=True)
class ChapterSpan:
    title: str
    start: int
    end: int
    pages: int
    source: str


@dataclass(frozen=True)
class SplitChunk:
    chunk_id: str
    source_path: Path
    chunk_path: Path
    document_type: str
    title: str
    chapter_titles: tuple[str, ...]
    estimated_pages: int
    start_index: int
    end_index: int


@dataclass(frozen=True)
class PreparedDocument:
    source_path: Path
    document_type: str
    total_pages: int | None
    used_semantic_split: bool
    chunks: tuple[SplitChunk, ...]
    notes: tuple[str, ...] = ()


@dataclass(frozen=True)
class _DocxTocEntry:
    title: str
    page_number: int | None


@dataclass
class _TocNode:
    """A node in the PDF bookmark hierarchy tree."""
    title: str
    start_page: int  # 0-based
    end_page: int = -1  # 0-based inclusive, filled later
    level: int = 1
    children: list["_TocNode"] | None = None

    def __post_init__(self):
        if self.children is None:
            self.children = []


class TOCSemanticSplitter:
    """Pre-processes PDFs and DOCXs so each MinerU request stays below 300 pages."""

    def __init__(
        self,
        *,
        max_pages_per_request: int = MAX_PAGES_PER_REQUEST,
        docx_words_per_page: int = DOCX_WORDS_PER_PAGE,
        docx_safe_max_pages_without_exact_count: int = DOCX_SAFE_MAX_PAGES_WITHOUT_EXACT_COUNT,
    ) -> None:
        if max_pages_per_request >= 300:
            raise ValueError("max_pages_per_request must be lower than 300")
        if max_pages_per_request <= 0:
            raise ValueError("max_pages_per_request must be positive")
        if docx_words_per_page <= 0:
            raise ValueError("docx_words_per_page must be positive")

        self.max_pages_per_request = max_pages_per_request
        self.docx_words_per_page = docx_words_per_page
        self.docx_safe_max_pages_without_exact_count = (
            docx_safe_max_pages_without_exact_count
        )

    def prepare_many(
        self,
        source_paths: Sequence[Path | str],
        workspace: Path | str,
    ) -> list[PreparedDocument]:
        workspace_path = Path(workspace).resolve()
        workspace_path.mkdir(parents=True, exist_ok=True)
        prepared: list[PreparedDocument] = []
        for source_path in source_paths:
            prepared.append(self.prepare_one(source_path, workspace_path))
        return prepared

    def prepare_one(
        self,
        source_path: Path | str,
        workspace: Path | str,
    ) -> PreparedDocument:
        path = Path(source_path).expanduser().resolve()
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"Input file does not exist: {path}")

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            raise ValueError(
                f"Unsupported file type: {path.name}. Supported suffixes: {sorted(SUPPORTED_SUFFIXES)}"
            )

        workspace_path = Path(workspace).resolve()
        workspace_path.mkdir(parents=True, exist_ok=True)

        if suffix == ".pdf":
            return self._prepare_pdf(path, workspace_path)
        return self._prepare_docx(path, workspace_path)

    def _prepare_pdf(self, path: Path, workspace: Path) -> PreparedDocument:
        total_pages = self._count_pdf_pages(path)
        notes: list[str] = []
        if total_pages <= self.max_pages_per_request:
            chunk = SplitChunk(
                chunk_id=f"{path.stem}-chunk-001",
                source_path=path,
                chunk_path=path,
                document_type="pdf",
                title=path.stem,
                chapter_titles=(path.stem,),
                estimated_pages=total_pages,
                start_index=0,
                end_index=total_pages,
            )
            return PreparedDocument(
                source_path=path,
                document_type="pdf",
                total_pages=total_pages,
                used_semantic_split=False,
                chunks=(chunk,),
            )

        # Strategy 1: Hierarchical bookmark-based splitting
        tree = self._build_pdf_toc_tree(path, total_pages)
        chapter_spans: list[ChapterSpan] = []
        if tree:
            chapter_spans = self._toc_tree_to_spans(tree)
            if chapter_spans:
                logger.info(
                    "Hierarchical bookmark split: {} chunks for {}",
                    len(chapter_spans), path.name,
                )

        # Strategy 2: Text pattern detection
        if not chapter_spans:
            chapter_spans = self._detect_pdf_text_patterns(path, total_pages)

        # Strategy 3: Font-size heuristic
        if not chapter_spans:
            chapter_spans = self._detect_pdf_font_size(path, total_pages)

        # Strategy 4: Fixed-size fallback
        if not chapter_spans:
            notes.append(
                "No bookmarks, text patterns, or font-size headings found; "
                "falling back to fixed-size page windows."
            )
            chapter_spans = self._build_uniform_page_spans(total_pages)

        consolidated = self._consolidate_pdf_spans(chapter_spans)
        working_dir = self._document_workspace(workspace, path)
        working_dir.mkdir(parents=True, exist_ok=True)

        chunks: list[SplitChunk] = []
        pdf_reader_cls, _ = _load_pypdf_classes()
        reader = pdf_reader_cls(str(path))
        for idx, span_group in enumerate(consolidated, start=1):
            start_page = span_group[0].start
            end_page = span_group[-1].end
            chunk_name = f"{self._safe_stem(path.stem)}.chunk_{idx:03d}.pdf"
            chunk_path = working_dir / chunk_name
            self._write_pdf_slice(reader, start_page, end_page, chunk_path)
            chunks.append(
                SplitChunk(
                    chunk_id=f"{path.stem}-chunk-{idx:03d}",
                    source_path=path,
                    chunk_path=chunk_path,
                    document_type="pdf",
                    title=self._span_group_title(span_group),
                    chapter_titles=tuple(span.title for span in span_group),
                    estimated_pages=end_page - start_page,
                    start_index=start_page,
                    end_index=end_page,
                )
            )

        return PreparedDocument(
            source_path=path,
            document_type="pdf",
            total_pages=total_pages,
            used_semantic_split=True,
            chunks=tuple(chunks),
            notes=tuple(notes),
        )

    def _prepare_docx(self, path: Path, workspace: Path) -> PreparedDocument:
        notes: list[str] = []
        document_factory = _load_docx_document_factory()
        document = document_factory(str(path))
        body_elements = list(document.element.body.iterchildren())

        chapter_spans, toc_entries = self._extract_docx_chapter_spans(document)
        estimated_total_pages_from_content = self._estimate_docx_pages_by_body(
            body_elements,
            0,
            len(body_elements),
        )

        exact_total_pages = self._estimate_docx_pages_via_libreoffice(path)
        effective_total_pages = exact_total_pages or estimated_total_pages_from_content

        if exact_total_pages is None:
            notes.append(
                "DOCX exact page count unavailable (LibreOffice/soffice not found); using conservative page estimation."
            )

        if not chapter_spans:
            notes.append(
                "DOCX TOC/headings were not detected. Falling back to pseudo-chapters based on paragraph windows."
            )
            chapter_spans = self._build_docx_pseudo_chapters(body_elements)

        # Enforce conservative threshold when page count is estimated instead of exact.
        effective_max = self.max_pages_per_request
        if exact_total_pages is None:
            effective_max = min(
                self.max_pages_per_request,
                self.docx_safe_max_pages_without_exact_count,
            )

        if effective_total_pages <= effective_max:
            chunk = SplitChunk(
                chunk_id=f"{path.stem}-chunk-001",
                source_path=path,
                chunk_path=path,
                document_type="docx",
                title=path.stem,
                chapter_titles=(path.stem,),
                estimated_pages=effective_total_pages,
                start_index=0,
                end_index=len(body_elements),
            )
            return PreparedDocument(
                source_path=path,
                document_type="docx",
                total_pages=exact_total_pages,
                used_semantic_split=False,
                chunks=(chunk,),
                notes=tuple(notes),
            )

        chapter_spans = self._apply_docx_page_numbers_if_available(
            chapter_spans,
            toc_entries,
            body_elements,
        )

        consolidated = self._consolidate_docx_spans(
            chapter_spans=chapter_spans,
            body_elements=body_elements,
            effective_max_pages=effective_max,
        )

        working_dir = self._document_workspace(workspace, path)
        working_dir.mkdir(parents=True, exist_ok=True)

        chunks: list[SplitChunk] = []
        for idx, span_group in enumerate(consolidated, start=1):
            start_index = span_group[0].start
            end_index = span_group[-1].end
            chunk_name = f"{self._safe_stem(path.stem)}.chunk_{idx:03d}.docx"
            chunk_path = working_dir / chunk_name
            self._write_docx_slice(path, body_elements, start_index, end_index, chunk_path)

            chunks.append(
                SplitChunk(
                    chunk_id=f"{path.stem}-chunk-{idx:03d}",
                    source_path=path,
                    chunk_path=chunk_path,
                    document_type="docx",
                    title=self._span_group_title(span_group),
                    chapter_titles=tuple(span.title for span in span_group),
                    estimated_pages=sum(span.pages for span in span_group),
                    start_index=start_index,
                    end_index=end_index,
                )
            )

        return PreparedDocument(
            source_path=path,
            document_type="docx",
            total_pages=exact_total_pages,
            used_semantic_split=True,
            chunks=tuple(chunks),
            notes=tuple(notes),
        )

    def _read_pdf_bookmarks(self, path: Path) -> list[BookmarkEntry]:
        entries = self._read_pdf_bookmarks_with_pymupdf(path)
        if entries:
            return entries
        return self._read_pdf_bookmarks_with_pypdf(path)

    def _read_pdf_bookmarks_with_pymupdf(self, path: Path) -> list[BookmarkEntry]:
        try:
            import fitz  # type: ignore
        except Exception:
            return []

        entries: list[BookmarkEntry] = []
        document = None
        try:
            document = fitz.open(str(path))
            raw_toc = document.get_toc(simple=True)
            for item in raw_toc:
                if len(item) < 3:
                    continue
                level, title, page_number = item[0], item[1], item[2]
                if not isinstance(page_number, int):
                    continue
                if page_number <= 0:
                    continue
                entries.append(
                    BookmarkEntry(
                        level=max(1, int(level)),
                        title=str(title).strip() or "Untitled",
                        page_index=page_number - 1,
                    )
                )
        except Exception as exc:
            logger.debug("Failed to read PDF bookmarks with PyMuPDF: {}", exc)
            return []
        finally:
            if document is not None:
                document.close()
        return entries

    def _read_pdf_bookmarks_with_pypdf(self, path: Path) -> list[BookmarkEntry]:
        pdf_reader_cls, _ = _load_pypdf_classes()
        reader = pdf_reader_cls(str(path))
        outline = getattr(reader, "outline", None)
        if not outline:
            return []

        entries: list[BookmarkEntry] = []

        def walk(items: Iterable, level: int) -> None:
            for item in items:
                if isinstance(item, list):
                    walk(item, level + 1)
                    continue
                try:
                    page_index = reader.get_destination_page_number(item)
                except Exception:
                    continue
                title = getattr(item, "title", "Untitled")
                entries.append(
                    BookmarkEntry(
                        level=max(1, level),
                        title=str(title).strip() or "Untitled",
                        page_index=int(page_index),
                    )
                )

        walk(outline, 1)
        return entries

    def _build_pdf_toc_tree(
        self,
        path: Path,
        total_pages: int,
    ) -> list[_TocNode]:
        """Build a hierarchical bookmark tree from PDF bookmarks.

        Returns root-level nodes.  Front matter (pages before the first
        bookmark) is merged into the first root node.
        """
        entries = self._read_pdf_bookmarks(path)
        if not entries:
            return []

        # De-duplicate consecutive entries pointing to the same page
        deduped: list[BookmarkEntry] = []
        last_page = -1
        for entry in entries:
            if entry.page_index == last_page:
                continue
            deduped.append(entry)
            last_page = entry.page_index

        if not deduped:
            return []

        # Build tree using a stack of ancestor nodes
        root_nodes: list[_TocNode] = []
        stack: list[_TocNode] = []

        for entry in deduped:
            node = _TocNode(
                title=entry.title,
                start_page=entry.page_index,
                level=entry.level,
            )
            # Pop until we find the parent (stack top level < current level)
            while stack and stack[-1].level >= entry.level:
                stack.pop()

            if stack:
                stack[-1].children.append(node)
            else:
                root_nodes.append(node)
            stack.append(node)

        # Assign end_page to each node
        self._compute_toc_end_pages(root_nodes, total_pages)

        # Merge front matter into the first root node
        if root_nodes and root_nodes[0].start_page > 0:
            root_nodes[0].start_page = 0

        return root_nodes

    def _compute_toc_end_pages(
        self,
        nodes: list[_TocNode],
        total_pages: int,
    ) -> None:
        """Assign end_page to each node (mutates in-place)."""
        for i, node in enumerate(nodes):
            if i + 1 < len(nodes):
                node.end_page = nodes[i + 1].start_page - 1
            else:
                node.end_page = total_pages - 1
            self._compute_toc_end_pages(node.children, total_pages)

    def _toc_tree_to_spans(
        self,
        nodes: list[_TocNode],
    ) -> list[ChapterSpan]:
        """Walk the TocNode tree to produce ChapterSpan objects.

        - Node within page limit -> emit as single span
        - Oversized node with children -> recurse into children
        - Oversized leaf node -> fixed-size split via _split_large_pdf_chapter
        - Trailing pages after last child are captured as a separate span
        """
        spans: list[ChapterSpan] = []

        for node in nodes:
            page_count = node.end_page - node.start_page + 1
            if page_count <= 0:
                continue

            if page_count <= self.max_pages_per_request:
                spans.append(ChapterSpan(
                    title=node.title,
                    start=node.start_page,
                    end=node.end_page + 1,  # exclusive end
                    pages=page_count,
                    source="bookmark",
                ))
            elif node.children:
                # Leading pages before first child
                first_child_start = node.children[0].start_page
                if first_child_start > node.start_page:
                    leading_count = first_child_start - node.start_page
                    leading_span = ChapterSpan(
                        title=f"{node.title} (intro)",
                        start=node.start_page,
                        end=first_child_start,  # exclusive
                        pages=leading_count,
                        source="bookmark-leading",
                    )
                    if leading_count <= self.max_pages_per_request:
                        spans.append(leading_span)
                    else:
                        for group in self._split_large_pdf_chapter(leading_span):
                            spans.extend(group)

                # Recurse into sub-chapters
                spans.extend(self._toc_tree_to_spans(node.children))

                # Capture trailing pages after last child
                last_child_end = node.children[-1].end_page
                if last_child_end < node.end_page:
                    trailing_count = node.end_page - last_child_end
                    trailing_span = ChapterSpan(
                        title=f"{node.title} (appendix)",
                        start=last_child_end + 1,
                        end=node.end_page + 1,  # exclusive end
                        pages=trailing_count,
                        source="bookmark-trailing",
                    )
                    if trailing_count <= self.max_pages_per_request:
                        spans.append(trailing_span)
                    else:
                        groups = self._split_large_pdf_chapter(trailing_span)
                        for group in groups:
                            spans.extend(group)
            else:
                # Oversized leaf -> fixed-size split
                oversized = ChapterSpan(
                    title=node.title,
                    start=node.start_page,
                    end=node.end_page + 1,  # exclusive end
                    pages=page_count,
                    source="bookmark-oversize",
                )
                groups = self._split_large_pdf_chapter(oversized)
                for group in groups:
                    spans.extend(group)

        return spans

    def _validate_pdf_heading(self, page_text_blocks: dict, match_text: str) -> bool:
        """Check if matched text uses a larger or bold font (heuristic)."""
        for block in page_text_blocks.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if match_text in text and len(text) > 3:
                        size = span.get("size", 0)
                        flags = span.get("flags", 0)
                        is_bold = bool(flags & (1 << 4))
                        if size >= 12 or (is_bold and size >= 10):
                            return True
        return False

    def _detect_pdf_text_patterns(
        self,
        path: Path,
        total_pages: int,
    ) -> list[ChapterSpan]:
        """Detect chapters by scanning PDF pages for heading text patterns.

        Returns ChapterSpan list, or empty list if nothing detected.
        """
        try:
            import fitz
        except ImportError:
            return []

        doc = None
        try:
            doc = fitz.open(str(path))
            boundaries: list[tuple[int, str]] = []
            seen_pages: set[int] = set()

            for page_idx in range(doc.page_count):
                page = doc[page_idx]
                text = page.get_text()
                if not text:
                    continue

                blocks = page.get_text("dict")

                for pattern in _PDF_HEADING_PATTERNS:
                    for match in pattern.finditer(text):
                        matched_line = match.group(0).strip()
                        if self._validate_pdf_heading(blocks, match.group(1)):
                            if page_idx not in seen_pages:
                                boundaries.append((page_idx, matched_line))
                                seen_pages.add(page_idx)
                            break  # one match per page is enough

            if not boundaries:
                return []

            # Build ChapterSpan objects
            spans: list[ChapterSpan] = []

            # Handle front matter (pages before first detected heading)
            if boundaries[0][0] > 0:
                first_page = boundaries[0][0]
                spans.append(ChapterSpan(
                    title="Front Matter",
                    start=0,
                    end=first_page,  # exclusive
                    pages=first_page,
                    source="text-pattern-front-matter",
                ))

            for i, (start_page, title) in enumerate(boundaries):
                end_page = (
                    boundaries[i + 1][0]
                    if i + 1 < len(boundaries)
                    else total_pages
                )
                spans.append(ChapterSpan(
                    title=title,
                    start=start_page,
                    end=end_page,  # exclusive
                    pages=end_page - start_page,
                    source="text-pattern",
                ))

            logger.info(
                "Detected {} chapters via text patterns in {}",
                len(spans), path.name,
            )
            return spans

        except Exception as exc:
            logger.debug("Text pattern detection failed for {}: {}", path.name, exc)
            return []
        finally:
            if doc is not None:
                doc.close()

    def _build_uniform_page_spans(self, total_pages: int) -> list[ChapterSpan]:
        spans: list[ChapterSpan] = []
        cursor = 0
        while cursor < total_pages:
            end = min(total_pages, cursor + self.max_pages_per_request)
            spans.append(
                ChapterSpan(
                    title=f"Pages {cursor + 1}-{end}",
                    start=cursor,
                    end=end,
                    pages=end - cursor,
                    source="fixed-window",
                )
            )
            cursor = end
        return spans

    def _consolidate_pdf_spans(
        self,
        chapter_spans: Sequence[ChapterSpan],
    ) -> list[list[ChapterSpan]]:
        grouped: list[list[ChapterSpan]] = []
        current_group: list[ChapterSpan] = []
        current_pages = 0

        for chapter in chapter_spans:
            if chapter.pages > self.max_pages_per_request:
                if current_group:
                    grouped.append(current_group)
                    current_group = []
                    current_pages = 0
                grouped.extend(self._split_large_pdf_chapter(chapter))
                continue

            if current_group and current_pages + chapter.pages > self.max_pages_per_request:
                grouped.append(current_group)
                current_group = [chapter]
                current_pages = chapter.pages
            else:
                current_group.append(chapter)
                current_pages += chapter.pages

        if current_group:
            grouped.append(current_group)
        return grouped

    def _split_large_pdf_chapter(self, chapter: ChapterSpan) -> list[list[ChapterSpan]]:
        groups: list[list[ChapterSpan]] = []
        cursor = chapter.start
        part = 1
        while cursor < chapter.end:
            end = min(chapter.end, cursor + self.max_pages_per_request)
            title = f"{chapter.title} (part {part})"
            groups.append(
                [
                    ChapterSpan(
                        title=title,
                        start=cursor,
                        end=end,
                        pages=end - cursor,
                        source="oversize-bookmark",
                    )
                ]
            )
            part += 1
            cursor = end
        return groups

    def _detect_pdf_font_size(
        self,
        path: Path,
        total_pages: int,
    ) -> list[ChapterSpan]:
        """Detect chapters using font-size anomalies (last resort).

        Estimates body text font size, then finds spans with significantly
        larger text as implicit chapter headings.
        """
        try:
            import fitz
        except ImportError:
            return []

        doc = None
        try:
            doc = fitz.open(str(path))

            # Estimate body font size by sampling pages
            sizes: list[float] = []
            sample_pages = min(20, doc.page_count)
            step = max(1, doc.page_count // sample_pages)
            for idx in range(0, doc.page_count, step):
                blocks = doc[idx].get_text("dict").get("blocks", [])
                for block in blocks:
                    if block.get("type") != 0:
                        continue
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            sizes.append(span.get("size", 0))

            if not sizes:
                return []

            body_size = float(Counter(sizes).most_common(1)[0][0])
            threshold = body_size * 1.4

            # Scan for anomalously large text
            boundaries: list[tuple[int, str]] = []
            seen_pages: set[int] = set()

            for page_idx in range(doc.page_count):
                blocks = doc[page_idx].get_text("dict").get("blocks", [])
                for block in blocks:
                    if block.get("type") != 0:
                        continue
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            size = span.get("size", 0)
                            if (
                                size >= threshold
                                and len(text) > 3
                                and page_idx not in seen_pages
                            ):
                                boundaries.append((page_idx, text))
                                seen_pages.add(page_idx)

            if not boundaries:
                return []

            # Build ChapterSpan objects
            spans: list[ChapterSpan] = []

            if boundaries[0][0] > 0:
                first_page = boundaries[0][0]
                spans.append(ChapterSpan(
                    title="Front Matter",
                    start=0,
                    end=first_page,
                    pages=first_page,
                    source="font-size-front-matter",
                ))

            for i, (start_page, title) in enumerate(boundaries):
                end_page = (
                    boundaries[i + 1][0]
                    if i + 1 < len(boundaries)
                    else total_pages
                )
                spans.append(ChapterSpan(
                    title=title,
                    start=start_page,
                    end=end_page,
                    pages=end_page - start_page,
                    source="font-size",
                ))

            logger.info(
                "Detected {} chapters via font-size heuristic in {}",
                len(spans), path.name,
            )
            return spans

        except Exception as exc:
            logger.debug("Font-size detection failed for {}: {}", path.name, exc)
            return []
        finally:
            if doc is not None:
                doc.close()

    def _write_pdf_slice(
        self,
        reader,
        start_page: int,
        end_page: int,
        output_path: Path,
    ) -> None:
        _, pdf_writer_cls = _load_pypdf_classes()
        writer = pdf_writer_cls()
        for page_index in range(start_page, end_page):
            page = reader.pages[page_index]
            if hasattr(page, "clone"):
                page = page.clone(
                    writer,
                    ignore_fields=_PDF_PAGE_CLONE_IGNORE_FIELDS,
                )
            writer.add_page(page)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with output_path.open("wb") as file_obj:
            writer.write(file_obj)

    def _extract_docx_chapter_spans(
        self,
        document,
    ) -> tuple[list[ChapterSpan], list[_DocxTocEntry]]:
        toc_entries = self._extract_docx_toc_entries(document)
        body_elements = list(document.element.body.iterchildren())
        paragraph_to_body_index = {
            id(element): index
            for index, element in enumerate(body_elements)
            if element.tag == W_P_TAG
        }

        heading_candidates = [
            paragraph
            for paragraph in document.paragraphs
            if self._is_docx_heading_level_1(paragraph)
            and paragraph.text
            and paragraph.text.strip()
        ]

        heading_indices: list[int] = []
        heading_titles: list[str] = []

        matched_from_toc = self._match_toc_entries_to_headings(toc_entries, heading_candidates)
        if matched_from_toc:
            for _, paragraph in matched_from_toc:
                body_index = paragraph_to_body_index.get(id(paragraph._p))
                if body_index is None:
                    continue
                heading_indices.append(body_index)
                heading_titles.append(paragraph.text.strip())
        else:
            for paragraph in heading_candidates:
                body_index = paragraph_to_body_index.get(id(paragraph._p))
                if body_index is None:
                    continue
                heading_indices.append(body_index)
                heading_titles.append(paragraph.text.strip())

        if not heading_indices:
            return [], toc_entries

        ordered = sorted(
            zip(heading_indices, heading_titles),
            key=lambda item: item[0],
        )

        deduped: list[tuple[int, str]] = []
        last_index = -1
        for index, title in ordered:
            if index <= last_index:
                continue
            deduped.append((index, title))
            last_index = index

        if not deduped:
            return [], toc_entries

        spans: list[ChapterSpan] = []
        if deduped[0][0] > 0:
            preface_pages = self._estimate_docx_pages_by_body(
                body_elements,
                0,
                deduped[0][0],
            )
            spans.append(
                ChapterSpan(
                    title="Front Matter",
                    start=0,
                    end=deduped[0][0],
                    pages=preface_pages,
                    source="docx-front-matter",
                )
            )

        for idx, (start_index, title) in enumerate(deduped):
            end_index = (
                deduped[idx + 1][0] if idx + 1 < len(deduped) else len(body_elements)
            )
            if end_index <= start_index:
                continue
            pages = self._estimate_docx_pages_by_body(body_elements, start_index, end_index)
            spans.append(
                ChapterSpan(
                    title=title,
                    start=start_index,
                    end=end_index,
                    pages=pages,
                    source="docx-heading",
                )
            )

        return spans, toc_entries

    def _extract_docx_toc_entries(self, document) -> list[_DocxTocEntry]:
        entries: list[_DocxTocEntry] = []
        for paragraph in document.paragraphs:
            style_name = ""
            if paragraph.style is not None and paragraph.style.name:
                style_name = paragraph.style.name.strip().lower()
            if not style_name.startswith("toc"):
                continue

            raw_text = paragraph.text.strip()
            if not raw_text:
                continue

            page_number = None
            title = raw_text
            match = _DOCX_TOC_LINE_PATTERN.match(raw_text)
            if match:
                title = match.group("title").strip()
                page_number = int(match.group("page"))

            entries.append(_DocxTocEntry(title=title, page_number=page_number))
        return entries

    def _match_toc_entries_to_headings(
        self,
        toc_entries: Sequence[_DocxTocEntry],
        headings: Sequence,
    ) -> list[tuple[_DocxTocEntry, object]]:
        if not toc_entries or not headings:
            return []

        matches: list[tuple[_DocxTocEntry, object]] = []
        heading_cursor = 0
        for toc_entry in toc_entries:
            target = self._normalize_title(toc_entry.title)
            if not target:
                continue

            for index in range(heading_cursor, len(headings)):
                heading = headings[index]
                heading_text = self._normalize_title(getattr(heading, "text", ""))
                if not heading_text:
                    continue
                if heading_text == target or heading_text in target or target in heading_text:
                    matches.append((toc_entry, heading))
                    heading_cursor = index + 1
                    break

        # Single accidental match is often not enough for stable chapter mapping.
        if len(matches) < 2:
            return []
        return matches

    def _build_docx_pseudo_chapters(self, body_elements: Sequence) -> list[ChapterSpan]:
        spans: list[ChapterSpan] = []
        target_words = self.max_pages_per_request * self.docx_words_per_page
        start = 0
        acc_words = 0
        for index in range(len(body_elements)):
            acc_words += self._estimate_words_in_body_element(body_elements[index])
            if acc_words < target_words:
                continue

            end = index + 1
            spans.append(
                ChapterSpan(
                    title=f"Section {len(spans) + 1}",
                    start=start,
                    end=end,
                    pages=self._estimate_docx_pages_by_body(body_elements, start, end),
                    source="docx-pseudo",
                )
            )
            start = end
            acc_words = 0

        if start < len(body_elements):
            spans.append(
                ChapterSpan(
                    title=f"Section {len(spans) + 1}",
                    start=start,
                    end=len(body_elements),
                    pages=self._estimate_docx_pages_by_body(body_elements, start, len(body_elements)),
                    source="docx-pseudo",
                )
            )
        return spans

    def _apply_docx_page_numbers_if_available(
        self,
        chapter_spans: Sequence[ChapterSpan],
        toc_entries: Sequence[_DocxTocEntry],
        body_elements: Sequence,
    ) -> list[ChapterSpan]:
        page_numbers = [entry.page_number for entry in toc_entries if entry.page_number is not None]
        if len(page_numbers) < 2:
            return list(chapter_spans)

        chapters = list(chapter_spans)
        number_index = 0
        updated: list[ChapterSpan] = []
        for chapter in chapters:
            if chapter.source == "docx-front-matter":
                updated.append(chapter)
                continue
            if number_index + 1 < len(page_numbers):
                current_page = page_numbers[number_index]
                next_page = page_numbers[number_index + 1]
                if next_page > current_page:
                    pages = max(1, next_page - current_page)
                    updated.append(
                        ChapterSpan(
                            title=chapter.title,
                            start=chapter.start,
                            end=chapter.end,
                            pages=pages,
                            source=f"{chapter.source}-toc-pages",
                        )
                    )
                    number_index += 1
                    continue
            updated.append(
                ChapterSpan(
                    title=chapter.title,
                    start=chapter.start,
                    end=chapter.end,
                    pages=self._estimate_docx_pages_by_body(body_elements, chapter.start, chapter.end),
                    source=chapter.source,
                )
            )
            number_index += 1
        return updated

    def _consolidate_docx_spans(
        self,
        chapter_spans: Sequence[ChapterSpan],
        body_elements: Sequence,
        effective_max_pages: int,
    ) -> list[list[ChapterSpan]]:
        grouped: list[list[ChapterSpan]] = []
        current_group: list[ChapterSpan] = []
        current_pages = 0

        for chapter in chapter_spans:
            chapter_pages = max(1, chapter.pages)
            if chapter_pages > effective_max_pages:
                if current_group:
                    grouped.append(current_group)
                    current_group = []
                    current_pages = 0
                grouped.extend(
                    self._split_large_docx_chapter(
                        chapter=chapter,
                        body_elements=body_elements,
                        effective_max_pages=effective_max_pages,
                    )
                )
                continue

            if current_group and current_pages + chapter_pages > effective_max_pages:
                grouped.append(current_group)
                current_group = [chapter]
                current_pages = chapter_pages
            else:
                current_group.append(chapter)
                current_pages += chapter_pages

        if current_group:
            grouped.append(current_group)

        return grouped

    def _split_large_docx_chapter(
        self,
        chapter: ChapterSpan,
        body_elements: Sequence,
        effective_max_pages: int,
    ) -> list[list[ChapterSpan]]:
        groups: list[list[ChapterSpan]] = []
        target_words = effective_max_pages * self.docx_words_per_page

        start = chapter.start
        acc_words = 0
        part_index = 1

        for body_index in range(chapter.start, chapter.end):
            acc_words += self._estimate_words_in_body_element(body_elements[body_index])
            if acc_words < target_words:
                continue

            end = body_index + 1
            pages = self._estimate_docx_pages_by_body(body_elements, start, end)
            groups.append(
                [
                    ChapterSpan(
                        title=f"{chapter.title} (part {part_index})",
                        start=start,
                        end=end,
                        pages=pages,
                        source="docx-oversize-chapter",
                    )
                ]
            )
            start = end
            acc_words = 0
            part_index += 1

        if start < chapter.end:
            groups.append(
                [
                    ChapterSpan(
                        title=f"{chapter.title} (part {part_index})",
                        start=start,
                        end=chapter.end,
                        pages=self._estimate_docx_pages_by_body(body_elements, start, chapter.end),
                        source="docx-oversize-chapter",
                    )
                ]
            )

        return groups

    def _write_docx_slice(
        self,
        source_path: Path,
        body_elements: Sequence,
        start_index: int,
        end_index: int,
        output_path: Path,
    ) -> None:
        document_factory = _load_docx_document_factory()
        chunk_doc = document_factory(str(source_path))
        chunk_body = chunk_doc.element.body

        for element in list(chunk_body.iterchildren()):
            chunk_body.remove(element)

        for index in range(start_index, end_index):
            chunk_body.append(deepcopy(body_elements[index]))

        # Keep section properties to avoid corrupt generated DOCX chunks.
        if not any(element.tag == W_SECTPR_TAG for element in chunk_body.iterchildren()):
            source_doc = document_factory(str(source_path))
            sect_pr = source_doc.element.body.sectPr
            if sect_pr is not None:
                chunk_body.append(deepcopy(sect_pr))

        output_path.parent.mkdir(parents=True, exist_ok=True)
        chunk_doc.save(str(output_path))

    def _estimate_docx_pages_by_body(
        self,
        body_elements: Sequence,
        start_index: int,
        end_index: int,
    ) -> int:
        words = 0
        for index in range(start_index, min(end_index, len(body_elements))):
            words += self._estimate_words_in_body_element(body_elements[index])
        return max(1, ceil(words / self.docx_words_per_page))

    def _estimate_words_in_body_element(self, element) -> int:
        text_fragments: list[str] = []
        for node in element.iter():
            if node.text and node.tag == W_T_TAG:
                text_fragments.append(node.text)
        if text_fragments:
            return max(1, sum(len(_WORD_PATTERN.findall(text)) for text in text_fragments))

        local_name = str(element.tag).rsplit("}", maxsplit=1)[-1]
        if local_name in {"tbl", "drawing", "pict", "object"}:
            return 80
        return 10

    def _estimate_docx_pages_via_libreoffice(self, path: Path) -> int | None:
        soffice_binary = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice_binary:
            return None

        with tempfile.TemporaryDirectory(prefix="mineru-docx-pages-") as temp_dir_name:
            temp_dir = Path(temp_dir_name)
            command = [
                soffice_binary,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(temp_dir),
                str(path),
            ]
            try:
                subprocess.run(
                    command,
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=180,
                )
            except Exception as exc:
                logger.debug("Failed to count DOCX pages via LibreOffice: {}", exc)
                return None

            generated_pdf = temp_dir / f"{path.stem}.pdf"
            if not generated_pdf.exists():
                candidates = sorted(temp_dir.glob("*.pdf"))
                if not candidates:
                    return None
                generated_pdf = candidates[0]

            try:
                return self._count_pdf_pages(generated_pdf)
            except Exception as exc:
                logger.debug("Failed to read temporary PDF page count: {}", exc)
                return None

    def _count_pdf_pages(self, path: Path) -> int:
        pdf_reader_cls, _ = _load_pypdf_classes()
        reader = pdf_reader_cls(str(path))
        return len(reader.pages)

    def _is_docx_heading_level_1(self, paragraph) -> bool:
        style_name = ""
        if paragraph.style is not None and paragraph.style.name:
            style_name = paragraph.style.name.strip().lower()
        return style_name in {"heading 1", "heading1", "title"}

    def _normalize_title(self, value: str) -> str:
        cleaned = re.sub(r"\s+", " ", value.strip().lower())
        cleaned = re.sub(r"[^\w\s]", "", cleaned)
        return cleaned.strip()

    def _span_group_title(self, span_group: Sequence[ChapterSpan]) -> str:
        if not span_group:
            return "Chunk"
        if len(span_group) == 1:
            return span_group[0].title
        return f"{span_group[0].title} -> {span_group[-1].title}"

    def _document_workspace(self, workspace: Path, source_path: Path) -> Path:
        digest = hashlib.sha1(str(source_path).encode("utf-8")).hexdigest()[:10]
        return workspace / f"{self._safe_stem(source_path.stem)}_{digest}"

    def _safe_stem(self, value: str) -> str:
        cleaned = re.sub(r"[^a-zA-Z0-9._-]", "_", value)
        cleaned = cleaned.strip("._")
        return cleaned or "document"
